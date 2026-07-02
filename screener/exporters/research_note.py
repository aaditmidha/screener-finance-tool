"""Institutional-style research note generator (Word .docx).

Mirrors the reference analyst notes (e.g. the POCL Nuvama meet note): a title, a
**focus-chart grid**, several thematic thesis sections, a **peer comparison**
table, and the **full financial statements** (income statement, balance sheet,
cash flow) plus a **ratios & returns** block. The thesis prose is written by the
LLM (Groq, via :mod:`screener.llm`) from the company's own figures and any
uploaded annual-report disclosures; every table and chart is computed from the
parsed data via :mod:`screener.exporters.financial_model`, so the note and the
Excel model always agree and nothing is fabricated.

The LLM client is injectable, so section generation is unit-tested without a
key, and the chart/table/docx builders are pure.
"""

import io
import json
import logging
import re
from dataclasses import dataclass, field
from typing import Any

from screener.exporters import financial_model as fm
from screener.llm import ChatClient, LLMError, chat
from screener.scraper.parser import CompanyFinancials, FinancialTable

logger = logging.getLogger(__name__)

_FENCE_RE = re.compile(r"^```(?:json)?\s*|\s*```$", re.MULTILINE)

# House-style palette (kept close to the reference notes; no em dashes anywhere).
_NAVY = (0x1F, 0x2A, 0x44)
_ACCENT = "#2f6bff"
_CORAL = "#f04e45"

_SECTIONS_SYSTEM = """You are a buy-side equity analyst writing a detailed \
management/research note for a portfolio manager. Using ONLY the data provided, \
return a JSON array of 7-9 sections, each {"heading": "...", "body": "..."}. \
Cover, in order: Company overview; Business and segments; Growth drivers and \
capacity; Margins and operating leverage; Balance sheet and cash flow quality; \
Returns and capital allocation; Peer positioning; Key risks; and a final \
"Valuation and view" section. Each body is 70-130 words, specific with the \
numbers given (cite the actual figures and growth rates), plain English, no \
markdown, no preamble, and NO em dashes. Never invent data not provided; if a \
topic has no data, write briefly from what is available."""

_SECTIONS_USER = """Write the note for {name} ({symbol}).

{context}

Return ONLY the JSON array of sections."""


@dataclass
class NoteSection:
    """One thesis section of the note."""

    heading: str
    body: str


@dataclass
class StatementTable:
    """A rendered financial-statement table (title, periods, derived rows)."""

    title: str
    periods: list[str]
    rows: list[fm.StatementRow] = field(default_factory=list)


@dataclass
class ResearchNote:
    """The assembled note content (before rendering)."""

    symbol: str
    name: str
    sections: list[NoteSection] = field(default_factory=list)
    periods: list[str] = field(default_factory=list)
    key_financials: list[tuple[str, list[float | None], str]] = field(default_factory=list)
    statements: list[StatementTable] = field(default_factory=list)
    charts: list[tuple[str, bytes]] = field(default_factory=list)
    peer_columns: list[str] = field(default_factory=list)
    peer_rows: list[list[Any]] = field(default_factory=list)


def _latest_row(table: FinancialTable | None, *needles: str) -> list[float | None] | None:
    if table is None:
        return None
    for needle in needles:
        row = table.row(needle)
        if row:
            return row
    return None


def _ratio_series(num: list[float | None] | None,
                  den: list[float | None] | None) -> list[float | None]:
    if num is None or den is None:
        return []
    return [
        (n / d) if n is not None and d not in (None, 0) else None
        for n, d in zip(num, den)
    ]


def key_financials(fin: CompanyFinancials) -> tuple[list[str], list[tuple[str, list, str]]]:
    """Compute the Key Financials table (label, per-period values, fmt).

    Args:
        fin: Parsed company financials.

    Returns:
        (period labels, rows) where each row is (label, values, fmt) with fmt in
        {"num", "pct"}; empty when there is no P&L.
    """
    pl = fin.profit_loss
    if pl is None:
        return [], []
    periods = pl.periods
    revenue = _latest_row(pl, "sales", "revenue")
    op_profit = _latest_row(pl, "operating profit")
    dep = _latest_row(pl, "depreciation")
    pat = _latest_row(pl, "net profit", "profit after tax")
    eps = _latest_row(pl, "eps")
    ebit = ([o - d if o is not None and d is not None else None
             for o, d in zip(op_profit, dep)] if op_profit and dep else None)

    rows: list[tuple[str, list, str]] = []
    def add(label, values, fmt):
        if values and any(v is not None for v in values):
            rows.append((label, list(values), fmt))

    add("Revenue", revenue, "num")
    add("EBITDA", op_profit, "num")
    add("EBITDA margin %", _ratio_series(op_profit, revenue), "pct")
    add("EBIT", ebit, "num")
    add("PAT", pat, "num")
    add("PAT margin %", _ratio_series(pat, revenue), "pct")
    add("EPS", eps, "num")
    return periods, rows


# --------------------------------------------------------------------------- #
# Charts
# --------------------------------------------------------------------------- #
def _png(fig) -> bytes:
    import matplotlib.pyplot as plt
    buf = io.BytesIO()
    fig.savefig(buf, format="png", dpi=130, bbox_inches="tight")
    plt.close(fig)
    return buf.getvalue()


def _nan(values: list[float | None]) -> list[float]:
    return [(v if v is not None else float("nan")) for v in values]


def focus_charts(periods: list[str],
                 key_rows: list[tuple[str, list, str]]) -> list[tuple[str, bytes]]:
    """Render the basic Revenue/PAT and margin charts from key-financials rows.

    Kept for the Excel Charts sheet and backwards compatibility; the richer note
    grid is :func:`chart_grid`.

    Args:
        periods: Period labels (x-axis).
        key_rows: Output of :func:`key_financials`.

    Returns:
        List of (chart title, PNG bytes); empty if nothing plottable.
    """
    import matplotlib
    matplotlib.use("Agg")
    import matplotlib.pyplot as plt

    by_label = {label: values for label, values, _fmt in key_rows}
    charts: list[tuple[str, bytes]] = []

    if "Revenue" in by_label or "PAT" in by_label:
        fig, ax = plt.subplots(figsize=(5.2, 2.6))
        x = range(len(periods))
        if "Revenue" in by_label:
            ax.bar(x, _nan(by_label["Revenue"]), color=_ACCENT, label="Revenue")
        if "PAT" in by_label:
            ax.plot(x, _nan(by_label["PAT"]), color=_CORAL, marker="o", label="PAT")
        ax.set_xticks(list(x)); ax.set_xticklabels(periods, rotation=45, ha="right", fontsize=7)
        ax.set_title("Revenue and PAT", fontsize=10); ax.legend(fontsize=7)
        ax.spines[["top", "right"]].set_visible(False)
        charts.append(("Revenue and PAT", _png(fig)))

    margin_labels = [lbl for lbl in ("EBITDA margin %", "PAT margin %") if lbl in by_label]
    if margin_labels:
        fig, ax = plt.subplots(figsize=(5.2, 2.6))
        x = range(len(periods))
        for lbl in margin_labels:
            ax.plot(x, [v * 100 if v is not None else float("nan") for v in by_label[lbl]],
                    marker="o", label=lbl.replace(" %", ""))
        ax.set_xticks(list(x)); ax.set_xticklabels(periods, rotation=45, ha="right", fontsize=7)
        ax.set_title("Margins (%)", fontsize=10); ax.legend(fontsize=7)
        ax.spines[["top", "right"]].set_visible(False)
        charts.append(("Margins", _png(fig)))

    return charts


def chart_grid(fin: CompanyFinancials) -> list[tuple[str, bytes]]:
    """Render the focus-chart grid for the note (revenue, margins, returns, growth).

    Args:
        fin: Parsed company financials.

    Returns:
        Up to four (title, PNG bytes) charts; empty if there are no periods.
    """
    import matplotlib
    matplotlib.use("Agg")
    import matplotlib.pyplot as plt

    periods = fm.periods(fin)
    if not periods:
        return []
    inc = {r.label: r.values for r in fm.income_statement(fin)}
    cs = {r.label: r.values for r in fm.common_size(fin)}
    gr = {r.label: r.values for r in fm.growth(fin)}
    rt = {r.label: r.values for r in fm.ratios(fin)}
    x = list(range(len(periods)))
    charts: list[tuple[str, bytes]] = []

    def _frame(ax, title):
        ax.set_xticks(x); ax.set_xticklabels(periods, rotation=45, ha="right", fontsize=7)
        ax.set_title(title, fontsize=10); ax.legend(fontsize=7)
        ax.spines[["top", "right"]].set_visible(False)

    if "Revenue from operations" in inc:
        fig, ax = plt.subplots(figsize=(5.0, 2.6))
        ax.bar(x, _nan(inc["Revenue from operations"]), color=_ACCENT, label="Revenue")
        if "Profit after tax" in inc:
            ax.plot(x, _nan(inc["Profit after tax"]), color=_CORAL, marker="o", label="PAT")
        _frame(ax, "Revenue and PAT (INR cr)")
        charts.append(("Revenue and PAT (INR cr)", _png(fig)))

    margin_rows = [(lbl, cs[lbl]) for lbl in ("EBITDA margin", "PAT margin") if lbl in cs]
    if margin_rows:
        fig, ax = plt.subplots(figsize=(5.0, 2.6))
        for lbl, vals in margin_rows:
            ax.plot(x, [v * 100 if v is not None else float("nan") for v in vals],
                    marker="o", label=lbl)
        _frame(ax, "Margin trend (%)")
        charts.append(("Margin trend (%)", _png(fig)))

    return_rows = [(lbl, rt[lbl]) for lbl in ("ROCE", "ROE") if lbl in rt]
    if return_rows:
        fig, ax = plt.subplots(figsize=(5.0, 2.6))
        for lbl, vals in return_rows:
            ax.plot(x, [v * 100 if v is not None else float("nan") for v in vals],
                    marker="o", label=lbl)
        _frame(ax, "Return ratios (%)")
        charts.append(("Return ratios (%)", _png(fig)))

    if "Revenue growth" in gr:
        fig, ax = plt.subplots(figsize=(5.0, 2.6))
        ax.bar(x, [v * 100 if v is not None else float("nan") for v in gr["Revenue growth"]],
               color=_ACCENT, label="Revenue growth")
        _frame(ax, "Revenue growth (% YoY)")
        charts.append(("Revenue growth (% YoY)", _png(fig)))

    return charts


# --------------------------------------------------------------------------- #
# Sections (LLM) + context
# --------------------------------------------------------------------------- #
def build_sections(name: str, symbol: str, context: str,
                   client: ChatClient | None = None) -> list[NoteSection]:
    """Generate the thesis sections via the LLM (regex-free JSON parse).

    Args:
        name: Company name.
        symbol: Ticker.
        context: Serialised data block for the model.
        client: Optional injected chat client.

    Returns:
        Parsed NoteSections. Empty list if the LLM is unavailable or the
        response can't be parsed (the note still renders with tables/charts).
    """
    try:
        raw = chat(_SECTIONS_SYSTEM,
                   _SECTIONS_USER.format(name=name, symbol=symbol, context=context),
                   client=client, temperature=0.3, max_tokens=2600)
        payload = json.loads(_FENCE_RE.sub("", raw).strip())
    except (LLMError, json.JSONDecodeError) as exc:
        logger.warning("Section generation unavailable (%s)", exc)
        return []
    sections = []
    for item in payload if isinstance(payload, list) else []:
        heading, body = item.get("heading"), item.get("body")
        if heading and body:
            sections.append(NoteSection(heading=str(heading), body=str(body)))
    return sections


def _qualitative_context(ar_rows: list | None) -> str:
    """Build a management-guidance / disclosed-risk block from AR-extracted rows."""
    if not ar_rows:
        return ""
    ordered = sorted(ar_rows, key=lambda r: getattr(r, "fiscal_year", 0))
    latest = ordered[-1]
    lines: list[str] = []
    guided = {
        "Guided revenue growth": getattr(latest, "guided_revenue_growth", None),
        "Guided margin": getattr(latest, "guided_margin", None),
    }
    guided = {k: v for k, v in guided.items() if v is not None}
    if guided:
        lines.append("## Management guidance (latest annual report)")
        for k, v in guided.items():
            lines.append(f"- {k}: {v}")
    risks: list[str] = []
    for row in ordered:
        raw = getattr(row, "key_risks", None)
        if not raw:
            continue
        try:
            for risk in json.loads(raw):
                if risk and risk not in risks:
                    risks.append(str(risk))
        except (json.JSONDecodeError, TypeError):
            continue
    if risks:
        lines.append("## Disclosed key risks")
        for risk in risks[:12]:
            lines.append(f"- {risk}")
    return "\n".join(lines)


def _serialise_context(fin: CompanyFinancials, metrics: dict[str, Any] | None,
                       ar_rows: list | None = None, forecast_result: Any = None) -> str:
    """Build the rich data block the LLM writes from."""
    periods, rows = key_financials(fin)
    lines = ["## Key financials (oldest to newest: " + ", ".join(periods) + ")"]
    for label, values, fmt in rows:
        rendered = ["-" if v is None else (f"{v*100:.1f}%" if fmt == "pct" else f"{v:,.0f}")
                    for v in values]
        lines.append(f"- {label}: {', '.join(rendered)}")

    for title, section_rows in (("Common-size (% of revenue)", fm.common_size(fin)),
                                ("Growth (% YoY)", fm.growth(fin)),
                                ("Ratios and returns", fm.ratios(fin))):
        body = [r for r in section_rows if r.kind != "header" and any(v is not None for v in r.values)]
        if body:
            lines.append(f"\n## {title}")
            for r in body:
                lines.append(f"- {r.label}: {', '.join(_fmt_stmt(v, r.kind) for v in r.values)}")

    if forecast_result is not None:
        a = forecast_result.assumptions
        lines.append(f"\n## Driver-based forecast ({', '.join(forecast_result.forecast_periods)})")
        lines.append(f"- Assumptions: revenue growth {a.revenue_growth*100:.1f}%/yr, "
                     f"EBITDA margin {a.ebitda_margin*100:.1f}%, tax rate {a.tax_rate*100:.1f}%")
        for r in forecast_result.rows:
            if r.label in ("Revenue from operations", "EBITDA", "Profit after tax", "EPS (INR)"):
                tail = r.values[forecast_result.n_history:]
                lines.append(f"- {r.label} (forecast): "
                             f"{', '.join(_fmt_stmt(v, r.kind) for v in tail)}")

    if metrics:
        lines.append("\n## Forensic / quality metrics")
        for k, v in metrics.items():
            lines.append(f"- {k}: {v}")

    qualitative = _qualitative_context(ar_rows)
    if qualitative:
        lines.append("\n" + qualitative)
    return "\n".join(lines)


def _statement_tables(fin: CompanyFinancials) -> list[StatementTable]:
    """Build the full statement tables (IS / BS / CF / ratios) for the note."""
    pl_periods = fin.profit_loss.periods if fin.profit_loss else []
    bs_periods = fin.balance_sheet.periods if fin.balance_sheet else []
    cf_periods = fin.cash_flow.periods if fin.cash_flow else []
    specs = [
        ("Income statement (INR cr)", pl_periods, fm.income_statement(fin)),
        ("Balance sheet (INR cr)", bs_periods, fm.balance_sheet(fin)),
        ("Cash flow (INR cr)", cf_periods, fm.cash_flow(fin)),
        ("Ratios and returns", pl_periods, fm.ratios(fin)),
    ]
    return [StatementTable(title, periods, rows) for title, periods, rows in specs if rows]


def generate(fin: CompanyFinancials, name: str, symbol: str,
             metrics: dict[str, Any] | None = None,
             peer_ranking: Any = None,
             client: ChatClient | None = None,
             ar_rows: list | None = None,
             assumptions: Any = None) -> ResearchNote:
    """Assemble the full research-note content for a company.

    Args:
        fin: Parsed company financials.
        name: Company name.
        symbol: Ticker.
        metrics: Optional forensic/quality metrics for the prose + context.
        peer_ranking: Optional ranked peer DataFrame.
        client: Optional injected chat client.
        ar_rows: Optional ARExtractedData rows → adds management guidance and
            disclosed risks to the LLM context for richer prose.
        assumptions: Optional ForecastAssumptions → drives the forecast table
            and feeds the forecast into the LLM context.

    Returns:
        A populated :class:`ResearchNote`.
    """
    from screener.models import forecast

    periods, rows = key_financials(fin)
    forecast_result = forecast.project(fin, assumptions)
    context = _serialise_context(fin, metrics, ar_rows, forecast_result)
    sections = build_sections(name, symbol, context, client=client)
    statements = _statement_tables(fin)
    if forecast_result is not None:
        statements.append(StatementTable(
            "P&L forecast (FY+1 onward, 'E' = estimate, INR cr)",
            forecast_result.periods, forecast_result.rows))
    note = ResearchNote(symbol=symbol, name=name, sections=sections,
                        periods=periods, key_financials=rows,
                        statements=statements, charts=chart_grid(fin))
    if peer_ranking is not None and not peer_ranking.empty:
        note.peer_columns = ["Symbol", *[str(c) for c in peer_ranking.columns]]
        note.peer_rows = [[idx, *list(r)] for idx, r in
                          zip(peer_ranking.index, peer_ranking.round(3).values.tolist())]
    logger.info("Built research note for %s: %d sections, %d fin rows, %d statements",
                symbol, len(sections), len(rows), len(note.statements))
    return note


# --------------------------------------------------------------------------- #
# Rendering
# --------------------------------------------------------------------------- #
def _fmt_cell(value: float | None, fmt: str) -> str:
    if value is None:
        return "-"
    return f"{value * 100:.1f}%" if fmt == "pct" else f"{value:,.0f}"


def _fmt_stmt(value: float | None, kind: str) -> str:
    """Format a value for a financial_model row by its kind."""
    if value is None:
        return "-"
    if kind == "pct":
        return f"{value * 100:.1f}%"
    if kind == "x":
        return f"{value:.1f}x"
    if kind == "days":
        return f"{value:.0f}"
    if kind in ("eps", "ratio"):
        return f"{value:,.2f}"
    return f"{value:,.0f}"


def _heading(doc: Any, text: str, level: int) -> None:
    """Add a navy heading."""
    from docx.shared import RGBColor
    h = doc.add_heading(text, level=level)
    if h.runs:
        h.runs[0].font.color.rgb = RGBColor(*_NAVY)


def _source_caption(doc: Any, text: str = "Source: Company filings, auto-generated.") -> None:
    from docx.shared import Pt
    para = doc.add_paragraph(text)
    run = para.runs[0]
    run.italic = True
    run.font.size = Pt(8)


def _add_statement_table(doc: Any, table: StatementTable) -> None:
    """Render one StatementTable as a formatted Word table."""
    from docx.shared import Pt
    _heading(doc, table.title, level=3)
    grid = doc.add_table(rows=1, cols=len(table.periods) + 1)
    grid.style = "Light Grid Accent 1"
    hdr = grid.rows[0].cells
    hdr[0].text = "Particulars"
    for i, period in enumerate(table.periods, start=1):
        hdr[i].text = period
    for row in table.rows:
        cells = grid.add_row().cells
        cells[0].text = row.label
        is_header = row.kind == "header"
        if row.bold or is_header:
            for run in cells[0].paragraphs[0].runs:
                run.bold = True
        if not is_header:
            for i, value in enumerate(row.values, start=1):
                if i < len(cells):
                    cells[i].text = _fmt_stmt(value, row.kind)
    _source_caption(doc)


def to_docx(note: ResearchNote) -> bytes:
    """Render a :class:`ResearchNote` to a .docx and return the bytes.

    Args:
        note: The assembled note content.

    Returns:
        The .docx file contents.
    """
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor

    doc = Document()
    title = doc.add_heading(f"{note.name} ({note.symbol})", level=0)
    if title.runs:
        title.runs[0].font.color.rgb = RGBColor(*_NAVY)
    subtitle = doc.add_paragraph(
        "Auto-generated research note. Figures in INR cr unless stated. "
        "Verify against source filings.")
    subtitle.runs[0].italic = True

    if note.charts:
        _heading(doc, "Focus charts", level=2)
        for chart_title, png in note.charts:
            caption = doc.add_paragraph(chart_title)
            caption.runs[0].bold = True
            doc.add_picture(io.BytesIO(png), width=Inches(5.2))
        _source_caption(doc)

    for section in note.sections:
        _heading(doc, section.heading, level=2)
        doc.add_paragraph(section.body)

    if note.peer_rows:
        _heading(doc, "Peer comparison", level=2)
        table = doc.add_table(rows=1, cols=len(note.peer_columns))
        table.style = "Light Grid Accent 1"
        for i, col in enumerate(note.peer_columns):
            table.rows[0].cells[i].text = str(col)
        for row in note.peer_rows:
            cells = table.add_row().cells
            for i, value in enumerate(row):
                if i < len(cells):
                    cells[i].text = str(value)
        _source_caption(doc)

    if note.statements:
        _heading(doc, "Financials", level=2)
        for table in note.statements:
            _add_statement_table(doc, table)

    doc.add_paragraph()
    disclaimer = doc.add_paragraph(
        "For research and education only. Not investment advice. Figures are "
        "auto-extracted and should be verified against source filings.")
    disclaimer.runs[0].font.size = Pt(8)

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()
